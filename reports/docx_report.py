# -*- coding: utf-8 -*-
"""
Word 报告生成（V3）：合并根因分析 + 合并部署上线 + 单场景部署
字体通用化：宋体(中文) + Times New Roman(英文)
"""
import os, json, datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ..core import config as C

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
STRATEGY_LABELS = {
    "none": "无策略",
    "light": "轻量策略",
    "standard": "标准策略",
    "major": "重大策略",
}


def _set_ea_font(run, cn=CN_FONT, en=EN_FONT):
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), cn)


def _set_font(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = EN_FONT
    font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), CN_FONT)


def _h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_ea_font(run)
    return h


def _meta(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    _set_ea_font(r)


def _table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                _set_ea_font(r)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = str(v)
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_ea_font(r)
    return t


def _load_json(path):
    if not os.path.exists(path):
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def write_rootcause_docx(scenario, rep, monitor):
    """单场景根因报告（保留兼容）"""
    path = os.path.join(C.REPORT_DIR, f"根因分析报告_{scenario}.docx")
    doc = Document()
    _set_font(doc)
    doc.add_heading(f"模型根因分析报告 — {scenario}", 0)
    _meta(doc, f"生成时间：{datetime.datetime.now():%Y-%m-%d %H:%M}")
    v = rep["verdict"]
    _h(doc, "一、综合判定结论", 1)
    doc.add_paragraph(f"推荐迭代策略：{STRATEGY_LABELS.get(v['strategy'], v['strategy'])}")
    _h(doc, "二、第一层诊断", 1)
    rows = []
    for k, it in rep.get("layer1", {}).items():
        val = it.get("value_pp", it.get("value", it.get("ks_gap", "—")))
        val = f"{val:.4f}" if isinstance(val, float) else val
        rows.append([k, val, it.get("level", "—"), it.get("note", "")])
    if rows:
        _table(doc, ["诊断项", "数值", "等级", "说明"], rows)
    _h(doc, "三、第二层归因", 1)
    rows = []
    for k, it in rep.get("layer2", {}).items():
        val = it.get("value", it.get("max_delta_iv", it.get("max_missing", "—")))
        val = f"{val:.4f}" if isinstance(val, float) else val
        rows.append([k, val, it.get("level", "—"), ""])
    if rows:
        _table(doc, ["归因项", "数值", "等级", "说明"], rows)
    llm = rep.get("llm_analysis", {})
    _h(doc, "四、AI Agent 四维归因（V2.1新增）", 1)
    if llm.get("status") != "success":
        doc.add_paragraph("LLM 调用未成功，以下保留原有规则根因分析结果。")
    dimensions = llm.get("dimension_analysis", {})
    labels = {"data": "数据", "feature": "特征", "model": "模型", "business": "业务"}
    rows = []
    for key in ("data", "feature", "model", "business"):
        item = dimensions.get(key, {})
        refs = "、".join(item.get("evidence_refs", [])) or "无"
        rows.append([
            labels[key], item.get("conclusion", "证据不足"),
            item.get("basis", "—"), f"{float(item.get('confidence', 0)):.2f}", refs,
        ])
    _table(doc, ["维度", "结论", "简短依据", "置信度", "证据引用"], rows)

    _h(doc, "五、首要根因与结构化策略（V2.1新增）", 1)
    primary = llm.get("primary_root_cause", {})
    strategy = llm.get("strategy_recommendation", {})
    _table(doc, ["项目", "结论", "置信度", "证据引用"], [
        ["首要根因", primary.get("conclusion", "证据不足"),
         f"{float(primary.get('confidence', 0)):.2f}",
         "、".join(primary.get("evidence_refs", [])) or "无"],
    ])
    strategy_name = strategy.get("display_name", STRATEGY_LABELS.get(strategy.get("type", "standard"), "标准策略"))
    doc.add_paragraph(f"建议策略：{strategy_name}（Agent 置信度：{float(strategy.get('confidence', 0)):.2f}）")
    doc.add_paragraph(f"置信度依据：{strategy.get('confidence_basis', '—')}")

    _h(doc, "六、证据化优化建议（V2.1新增）", 1)
    recommendations = llm.get("recommendations", [])
    if recommendations:
        rows = []
        for item in recommendations:
            rows.append([
                item.get("action", ""), item.get("reason", ""),
                f"{float(item.get('confidence', 0)):.2f}",
                "、".join(item.get("evidence_refs", [])) or "无",
            ])
        _table(doc, ["建议动作", "原因", "置信度", "证据引用"], rows)
    else:
        doc.add_paragraph("无 LLM 优化建议。")
    gaps = llm.get("evidence_gaps", [])
    if gaps:
        doc.add_paragraph("证据不足项：" + "；".join(str(item) for item in gaps))

    _h(doc, "七、调用追踪信息（V2.1新增）", 1)
    trace = llm.get("trace", {})
    _table(doc, ["字段", "内容"], [
        ["Trace ID", trace.get("trace_id", "—")],
        ["时间", trace.get("timestamp", "—")],
        ["模型", llm.get("model", "—")],
        ["Prompt版本", trace.get("prompt_version", "—")],
        ["Thinking模式", trace.get("thinking_mode", "—")],
        ["调用状态", llm.get("status", "—")],
    ])
    doc.save(path)
    return path


def write_merged_rootcause_docx():
    """合并根因分析报告（所有场景整合为1个docx，含none场景）
    V4：轮次标签 = 第一轮2026-01 / 第二轮2026-02 / 第三轮2026-03（OOT只有一轮=2026-01）"""
    path = os.path.join(C.REPORT_DIR, "根因分析报告_合并.docx")
    round_label = {1: "第一轮（2026-01）", 2: "第二轮（2026-02）", 3: "第三轮（2026-03）"}
    doc = Document()
    _set_font(doc)
    doc.add_heading("模型根因分析报告（全场景合并）", 0)
    _meta(doc, f"生成时间：{datetime.datetime.now():%Y-%m-%d %H:%M}")
    for idx, sc in enumerate(C.ALL_SCENARIOS):
        spec = C.SCENARIO_SPEC.get(sc, {})
        _h(doc, f"{idx+1}. 场景 {sc}：{spec.get('desc','')}", 1)
        found = False
        # OOT 只有 r1（2026-01）；时序场景三轮
        max_r = 2 if sc == "oot_2026_01" else 4
        for r in range(1, max_r):
            rep = _load_json(os.path.join(C.REPORT_DIR, f"rootcause_report_{sc}_r{r}.json"))
            if rep is None:
                continue
            found = True
            v = rep.get("verdict", {})
            _h(doc, f"{round_label.get(r, f'第{r}轮')}", 2)
            doc.add_paragraph(f"HIGH：{('、'.join(v.get('high_items',[]))) or '无'}")
            doc.add_paragraph(f"MEDIUM：{('、'.join(v.get('medium_items',[]))) or '无'}")
            strat = v.get("strategy", "none")
            doc.add_paragraph(f"推荐策略：{strat}")
            rows = []
            for k, it in rep.get("layer1", {}).items():
                val = it.get("value_pp", it.get("value", it.get("ks_gap", "—")))
                val = f"{val:.4f}" if isinstance(val, float) else val
                rows.append([k, val, it.get("level", "—")])
            if rows:
                _table(doc, ["诊断项", "数值", "等级"], rows)
            rows = []
            for k, it in rep.get("layer2", {}).items():
                val = it.get("value", it.get("max_delta_iv", it.get("max_missing", "—")))
                val = f"{val:.4f}" if isinstance(val, float) else val
                rows.append([k, val, it.get("level", "—")])
            if rows:
                _table(doc, ["归因项", "数值", "等级"], rows)
            doc.add_paragraph()
        if not found:
            _h(doc, "该场景无根因分析", 2)
            doc.add_paragraph(f"场景{sc}三轮均未触发告警，策略为none，不迭代。")
            doc.add_paragraph()
    doc.save(path)
    return path


def write_merged_deploy_docx():
    """合并部署上线报告（所有场景整合为1个docx，含none场景）
    V4：轮次标签 = 第一轮2026-01 / 第二轮2026-02 / 第三轮2026-03（OOT只有一轮=2026-01）"""
    path = os.path.join(C.REPORT_DIR, "部署上线报告_合并.docx")
    round_label = {1: "第一轮（2026-01）", 2: "第二轮（2026-02）", 3: "第三轮（2026-03）"}
    doc = Document()
    _set_font(doc)
    doc.add_heading("模型部署上线评审报告（全场景合并）", 0)
    _meta(doc, f"生成时间：{datetime.datetime.now():%Y-%m-%d %H:%M}")
    dec_map = {
        "deploy_significant": "全量上线（候选显著优于现役）",
        "deploy_non_inferior": "全量上线（候选与现役非劣）",
        "hold": "暂缓上线（候选劣于现役）",
        "rollback": "回滚（灰度护栏未通过）",
        "no_model": "无候选模型（策略none，未迭代）",
    }
    for idx, sc in enumerate(C.ALL_SCENARIOS):
        spec = C.SCENARIO_SPEC.get(sc, {})
        _h(doc, f"{idx+1}. 场景 {sc}：{spec.get('desc','')}", 1)
        found = False
        # OOT 只有 r1（2026-01）；时序场景三轮
        max_r = 2 if sc == "oot_2026_01" else 4
        for r in range(1, max_r):
            dep = _load_json(os.path.join(C.REPORT_DIR, f"deploy_report_{sc}_r{r}.json"))
            if dep is None:
                continue
            found = True
            _h(doc, f"{round_label.get(r, f'第{r}轮')}", 2)
            dec_text = dec_map.get(dep.get("decision_code"), dep.get("final_decision", "—"))
            rows = [
                ["评审结论", dec_text],
                ["候选模型", f"{dep.get('candidate_name','—')}（{dep.get('candidate_algo','—')}）"],
                ["现役模型", f"{dep.get('champion_name','—')}"],
                ["评测集", dep.get("eval_set", "—")],
                ["评审后现役", dep.get("champion_after", "—")],
            ]
            _table(doc, ["项目", "内容"], rows)
            doc.add_paragraph()
            # 灰度验证
            stage_name = [("L1_OOT", "离线回放"), ("L2_20pct", "灰度20%"),
                          ("L2_50pct", "灰度50%"), ("FULL_100pct", "全量")]
            rows = []
            for key, label in stage_name:
                st = dep.get("stages", {}).get(key)
                if not st:
                    continue
                rows.append([label, st.get("n","—"),
                             f"{st.get('new',{}).get('KS',0):.4f}",
                             f"{st.get('champion',{}).get('KS',0):.4f}",
                             f"{st.get('new',{}).get('AUC',0):.4f}",
                             f"{st.get('champion',{}).get('AUC',0):.4f}",
                             "通过" if st.get("pass") else "未通过"])
            if rows:
                _table(doc, ["层级","样本量","候选KS","基座KS","候选AUC","基座AUC","结果"], rows)
            doc.add_paragraph()
        if not found:
            _h(doc, "该场景无部署评审（策略none，未迭代）", 2)
            doc.add_paragraph(f"场景{sc}三轮均未触发告警，不产出候选模型，不进入部署评审。现役模型保持不变。")
            doc.add_paragraph()
    doc.save(path)
    return path


def write_deploy_docx(scenario, dep, round_no=None):
    """单场景部署报告（V1流程 L1→L2→全量）"""
    suffix = f"_r{round_no}" if round_no is not None else ""
    path = os.path.join(C.REPORT_DIR, f"部署上线报告_{scenario}{suffix}.docx")
    doc = Document()
    _set_font(doc)
    doc.add_heading("模型部署上线评审报告", 0)
    _meta(doc, f"场景：{scenario}    生成时间：{datetime.datetime.now():%Y-%m-%d %H:%M}")
    _h(doc, "一、上线结论", 1)
    dec_map = {
        "deploy_significant": "全量上线（候选显著优于现役）",
        "deploy_non_inferior": "全量上线（候选与现役非劣）",
        "hold": "暂缓上线（候选劣于现役）",
        "rollback": "回滚（灰度护栏未通过）",
        "no_model": "无候选模型（策略none，未迭代）",
    }
    dec_text = dec_map.get(dep.get("decision_code"), dep["final_decision"])
    rows = [
        ["评审结论", dec_text],
        ["候选模型", f"{dep.get('candidate_name','—')}（{dep.get('candidate_algo','—')}）"],
        ["现役模型", f"{dep.get('champion_name','—')}"],
        ["评测集", dep.get("eval_set", "—")],
        ["评审后现役", dep.get("champion_after", "—")],
    ]
    _table(doc, ["项目", "内容"], rows)
    doc.add_paragraph()
    _h(doc, "二、灰度放量验证", 1)
    stage_name = [("L1_OOT", "离线全量回放"), ("L2_20pct", "灰度20%"),
                  ("L2_50pct", "灰度50%"), ("FULL_100pct", "全量评估")]
    rows = []
    for key, label in stage_name:
        st = dep.get("stages", {}).get(key)
        if not st:
            continue
        rows.append([label, st.get("n","—"),
                     f"{st.get('new',{}).get('KS',0):.4f}",
                     f"{st.get('champion',{}).get('KS',0):.4f}",
                     f"{st.get('new',{}).get('AUC',0):.4f}",
                     f"{st.get('champion',{}).get('AUC',0):.4f}",
                     "通过" if st.get("pass") else "未通过"])
    if rows:
        _table(doc, ["层级","样本量","候选KS","基座KS","候选AUC","基座AUC","结果"], rows)
    doc.add_paragraph()
    _h(doc, "三、业务无感切换", 1)
    doc.add_paragraph("双跑机制：候选与现役对同批进件并行打分，现役全程承载线上决策。")
    doc.add_paragraph("回退保障：任一护栏触发即秒级回退至现役，候选归档。")
    doc.add_paragraph("归档管理：评审过程数据落库备查，满足模型风险管理审计要求。")
    doc.save(path)
    return path
